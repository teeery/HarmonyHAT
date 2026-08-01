from pathlib import Path
from docx import Document

root = Path('.')
for path in sorted(root.glob('04-团队调研材料/**/*.docx')):
    print(f'\n===== {path} =====')
    doc = Document(path)
    for p in doc.paragraphs:
        text = p.text.strip()
        if text:
            print(text)
    for ti, table in enumerate(doc.tables, 1):
        print(f'[TABLE {ti}]')
        for row in table.rows:
            print(' | '.join(cell.text.replace('\n', ' / ').strip() for cell in row.cells))
